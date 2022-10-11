import docx
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

FOLDER = '/Documents/CATALOG/' 
COMPANY = 'COMPANY' # Will be written in the header of the document
COL_NUM = 4 # Number of columns that the table should have.
W = Inches(1.5) # Width of Image
H = Inches(1.8) # Height of Image

def create_doc():
    doc = docx.Document()
    return doc
def new_section(doc):
    section = doc.sections[0]
    section.right_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run()
    font = run.font
    font.name = 'Alef'
    font.size = Pt(15)
    run.text = COMPANY
    return doc

def break_page(doc):
    paragraph = doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    return doc

def insert_pics(table, files):
    c=0
    for row in range(row_num):
        for col in range(COL_NUM):
            file, file_ext = os.path.splitext(files[c])
            file_path = f'{os.environ["USERPROFILE"]}{FOLDER}{folder}/{files[c]}'
            cell = table.cell(row, col)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(file_path, width = W, height = H)
            run = paragraph.add_run()
            font = run.font
            font.size = Pt(11)
            run.add_break(WD_BREAK.LINE)
            run.text = file
            c+=1
            if c == len(files):
                c=0
                break


doc = create_doc()
doc = new_section(doc)
folders = os.listdir(f'{os.environ["USERPROFILE"]}{FOLDER}')
i=0
for folder in folders:
        i+=1
        path = f'{os.environ["USERPROFILE"]}{FOLDER}{folder}'
        if os.path.isdir(path):
            if not folder.startswith('.'):
                for root, dirs, files in os.walk(path):
                    if i > 1:
                        break_page(doc)
                    doc.add_heading(folder, level=0)
                    if len(files) == COL_NUM:
                        row_num = (len(files)//COL_NUM)
                    else:
                        row_num = (len(files)//COL_NUM)+1
                    t = doc.add_table(row_num, COL_NUM)
                    insert_pics(t, files)

doc.save('Catalog_Book.docx')